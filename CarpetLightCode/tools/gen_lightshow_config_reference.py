#!/usr/bin/env python3
"""Generates a one-page printable reference (lightshow_config_reference.txt
and .docx) containing ONLY the "Light show order" and "Config menu
structure" sections -- the same content as those 2 sections in README.md,
kept in sync by hand (no shared source with the README's own markdown).
This script IS the source of truth for that content -- edit SECTIONS below
(not the generated files) whenever a firmware change affects show/variation
order, the config menu, or pot/encoder behavior, then re-run this script and
mirror the same edit into README.md's "Light show order"/"Config menu
structure" sections.

Usage: python3 tools/gen_lightshow_config_reference.py (run from the
CarpetLightCode repo root, or anywhere -- paths below are relative to this
script's own location)

Requires: python-docx (pip install python-docx)
"""
import os
from datetime import datetime
from zoneinfo import ZoneInfo

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit("Missing dependency: pip install python-docx")

REPO_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TXT_PATH = os.path.join(REPO_DIR, "lightshow_config_reference.txt")
DOCX_PATH = os.path.join(REPO_DIR, "lightshow_config_reference.docx")

TITLE = "CARPET LIGHT — SHOW ORDER & CONFIG MENU REFERENCE"

DIR_CAVEAT = (
    "Enc right/left below is per the code's own convention -- NOT yet "
    "verified against real wiring (flip if backwards)."
)

# (section_title, intro_line_or_None, [(item_line, [sub_bullet_lines])])
SECTIONS = [
    ("LIGHT SHOW ORDER", (
        "S advances 1->2->3->4->5->1. Enc right/left advances/retreats "
        "the show's own variation, wrapping. " + DIR_CAVEAT
    ), [
        ("Global (any show, any variation)", [
            "S next show -- M no effect -- L enter config -- XL lights on/off",
            "D UV on/off (Equalizer: triple-strobe toggle instead)",
            "Enc right next variation, Enc left previous",
        ]),
        ("1 Nightrider (3 variations)", [
            "ManualHue -- Pot sets hue directly",
            "AutoHueCycle -- Pot sets flood auto-hue rate, 0.25-2Hz",
            "AutoWithSound -- Pot sets flood auto-hue rate; china flashes on bass hits",
        ]),
        ("2 Flame (2 variations)", [
            "Waterflames -- Pot sets fire-sim speed",
            "Flames -- Pot sets fire-sim speed",
        ]),
        ("3 Equalizer (3 variations)", [
            "VU meter -- Pot sets live peak threshold; D toggles triple-strobe",
            "new_standard -- Pot sets cycle length, 2-128 bass beats (log); D toggles triple-strobe",
            "pixel_war -- Pot sets cycle length, 2-128 bass beats (log); D toggles triple-strobe",
        ]),
        ("4 SpeedStripes (2 variations)", [
            "default -- Pot has no effect",
            "zebra -- Pot sets black-stripe width, 0-40ft",
        ]),
        ("5 Lighthouse (2 variations)", [
            "default -- Pot sets rotation-speed ceiling; china bass-hit white strobe",
            "No Strobe -- Pot sets rotation-speed ceiling; no china strobe",
        ]),
    ]),
    ("CONFIG MENU STRUCTURE", (
        "L enters config (always @ ① Brightness). M advances screen, "
        "wraps ①->②->③->①. S advances subsetting. D saves+exits. "
        "L cancels, no save. " + DIR_CAVEAT
    ), [
        ("① Brightness (3 sub) -- Pot adjusts each live value", [
            "Global -- 0-100%",
            "Headlight -- 50-100%",
            "China -- 0-100%",
        ]),
        ("② Audio (6 sub)", [
            "NoiseFloor/PeakThreshold -- S swaps which Pot adjusts",
            "HitDecay -- Pot sets 0-1000ms",
            "Foresight -- Pot 0-1000ms; Enc R/L next/prev hit-predict style (off/exp/machine-gun/drum, wraps)",
            "AGC -- Enc R advances Off->Band->Full, Enc L back toward Off (clamped); Pot n/a",
            "SndReact -- Enc R = on, Enc L = off (direct pick); Pot n/a",
            "AutoPk -- Enc R advances Off->Full->Bin, Enc L back toward Off (clamped); Pot n/a",
        ]),
        ("③ PowerTest (3 sub) -- Enc adjusts each; Pot n/a", [
            "Hue",
            "Saturation",
            "Brightness",
        ]),
    ]),
]

PAGE_W_IN, PAGE_H_IN = 8.5, 11.0
MARGIN_IN = 0.5
NUM_COLUMNS = 2  # guarantees one-page fit at readable size -- see set_columns()
BODY_FONT_PT = 10
HEADER_FONT_PT = 12
ITEM_FONT_PT = 10


def now_pacific_str():
    return datetime.now(ZoneInfo("America/Los_Angeles")).strftime("%Y-%m-%d %I:%M %p %Z")


def build_txt():
    lines = [TITLE, ""]
    for title, intro, items in SECTIONS:
        lines.append(title)
        if intro:
            lines.append("  " + intro)
        for item_line, subs in items:
            lines.append("  - " + item_line)
            for s in subs:
                lines.append("      - " + s)
        lines.append("")
    lines.append("-----")
    lines.append(f"Generated: {now_pacific_str()} (Pacific)")
    lines.append(
        "Source of truth: tools/gen_lightshow_config_reference.py -- edit "
        "SECTIONS there and re-run, mirroring the same edit into README.md's "
        "\"Light show order\"/\"Config menu structure\" sections."
    )
    return "\n".join(lines) + "\n"


def set_columns(section, num_cols, space_in=0.3):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(space_in * 1440)))
    sectPr.append(cols)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(PAGE_W_IN)
    section.page_height = Inches(PAGE_H_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    set_columns(section, NUM_COLUMNS)  # 2-column flow -- guarantees this fits one page

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(BODY_FONT_PT)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(15)

    for title, intro, items in SECTIONS:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(8)
        hp.paragraph_format.space_after = Pt(2)
        hrun = hp.add_run(title)
        hrun.bold = True
        hrun.font.size = Pt(HEADER_FONT_PT)

        if intro:
            ip = doc.add_paragraph()
            ip.paragraph_format.space_after = Pt(4)
            ip.paragraph_format.left_indent = Inches(0.15)
            irun = ip.add_run(intro)
            irun.italic = True
            irun.font.size = Pt(9)

        for item_line, subs in items:
            ip = doc.add_paragraph(style='List Bullet')
            ip.paragraph_format.space_after = Pt(0)
            irun = ip.add_run(item_line)
            irun.bold = True
            irun.font.size = Pt(ITEM_FONT_PT)
            for s in subs:
                sp = doc.add_paragraph(style='List Bullet 2')
                sp.paragraph_format.space_after = Pt(0)
                srun = sp.add_run(s)
                srun.font.size = Pt(ITEM_FONT_PT)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(10)
    frun = foot.add_run(
        f"Generated {now_pacific_str()} (Pacific) by "
        f"tools/gen_lightshow_config_reference.py"
    )
    frun.font.size = Pt(7)
    frun.italic = True

    doc.save(DOCX_PATH)


def main():
    with open(TXT_PATH, 'w') as f:
        f.write(build_txt())
    print(f"Wrote {TXT_PATH}")
    build_docx()
    print(f"Wrote {DOCX_PATH}")


if __name__ == '__main__':
    main()
